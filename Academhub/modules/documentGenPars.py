import os
import sys
import django

from django.utils.translation.trans_real import parse_accept_lang_header
from docx.enum.table import WD_ALIGN_VERTICAL
from openpyxl import Workbook
from openpyxl.styles import Color, PatternFill

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.style import WD_STYLE_TYPE

# Определение среды Django для тестов
sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), "../..")))
os.environ.setdefault("DJANGO_SETTINGS_MODULE", "Academhub.settings")
django.setup()
from Academhub.models import ContingentMovement, GradebookStudents, GroupStudents, Qualification, Specialty, Student, Gradebook


class GroupTableGenerator:
    """
    Класс для генерации таблицы групп
    """ 
    def __init__(self, groups) -> None:
        """
        Принимает на вход список групп

        groups: list[GroupStudents]
        """
        self.groups = groups

    def generate_document(self, path: str) -> None:
        """
        Генерирует новый документ по указанному пути
        """
        workbook = Workbook()
        worksheet = workbook.active
        if worksheet is not None:
            worksheet.title = "Группы"
            # Заголовки
            worksheet["A1"] = "№"
            worksheet["B1"] = "1 курс (9 кл.)"
            worksheet["C1"] = "2 курс (9 кл.)"
            worksheet["D1"] = "1 курс (11 кл.)"
            worksheet["E1"] = "3 курс (9 кл.)"
            worksheet["F1"] = "2 курс (11 кл.)"
            worksheet["G1"] = "4 курс (9 кл.)"
            worksheet["H1"] = "3 курс (11 кл.)"

            # Ширина ячеек
            worksheet.column_dimensions['A'].width = len(str(len(self.groups))) + 3
            GROUPNAME_CELL_WIDTH = max([len(i.full_name) for i in self.groups])
            worksheet.column_dimensions['B'].width = GROUPNAME_CELL_WIDTH
            worksheet.column_dimensions['C'].width = GROUPNAME_CELL_WIDTH
            worksheet.column_dimensions['D'].width = GROUPNAME_CELL_WIDTH
            worksheet.column_dimensions['E'].width = GROUPNAME_CELL_WIDTH
            worksheet.column_dimensions['F'].width = GROUPNAME_CELL_WIDTH
            worksheet.column_dimensions['G'].width = GROUPNAME_CELL_WIDTH
            worksheet.column_dimensions['H'].width = GROUPNAME_CELL_WIDTH

            for i in range(len(self.groups)):
                column_letter = 0
                # match self.groups[i].current_course:
                #     case 1:
                #         if self.groups[i].education_base == "Основное общее":
                #             column_letter = 2
                #         elif self.groups[i].education_base == "Среднее общее":
                #             column_letter = 4
                #     case 2:
                #         if self.groups[i].education_base == "Основное общее":
                #             column_letter = 3
                #         elif self.groups[i].education_base == "Среднее общее":
                #             column_letter = 6
                #     case 3:
                #         if self.groups[i].education_base == "Основное общее":
                #             column_letter = 5
                #         elif self.groups[i].education_base == "Среднее общее":
                #             column_letter = 8
                #     case 4:
                #         column_letter = 7

                # match self.groups[i].current_course:
                if self.groups[i].current_course == 1:
                    if self.groups[i].education_base == "Основное общее":
                        column_letter = 2
                    elif self.groups[i].education_base == "Среднее общее":
                        column_letter = 4
                elif self.groups[i].current_course == 2:
                    if self.groups[i].education_base == "Основное общее":
                        column_letter = 3
                    elif self.groups[i].education_base == "Среднее общее":
                        column_letter = 6
                elif self.groups[i].current_course == 3:
                    if self.groups[i].education_base == "Основное общее":
                        column_letter = 5
                    elif self.groups[i].education_base == "Среднее общее":
                        column_letter = 8
                elif self.groups[i].current_course == 4:
                    column_letter = 7
                
                last_row = worksheet.max_row
                while last_row > 0 and worksheet.cell(row=last_row, column=column_letter).value is None:
                    last_row -= 1
                
                worksheet.cell(row=last_row+1, column=column_letter).value = self.groups[i].full_name

            for i in range(2, worksheet.max_row+1):
                worksheet.cell(row=i, column=1).value = i - 1
            
            workbook.save(path)

# class CourseDiffereeError(Exception):
#     def __init__(self, message="Для формирования таблицы, студенты должны быть с одного курса"):
#         self.message = message
#         super().__init__(self.message)
#
# class EducationBaseDifferenceError(Exception):
#     def __init__(self, message="Для формирования таблицы, студенты должны быть с одинаковой базой образования"):
#         self.message = message
#         super().__init__(self.message)


class CourseTableGenerator:
    """
    Класс для генерации таблицы курса
    """
    def __init__(self, students, course="1", education_base="Основное общее") -> None:
        """
        Принимает на вход список со студентами
        students: list[Student]
        """


        self.students = students.filter(group__current_course=course).filter(group__education_base=education_base)
        self.course = course

        if education_base == "Основное общее":
            self.education_base = "9 кл."
        elif education_base == "Среднее общее":
            self.education_base = "11 кл."

    def generate_document(self, path):
        workbook = Workbook()
        worksheet = workbook.active

        if worksheet is not None:
            worksheet.title = f"{self.course} курс ({self.education_base})"

            # Заголовки
            worksheet["A1"] = "№"
            worksheet["B1"] = "ФИО"
            worksheet["C1"] = "Дата рождения"
            worksheet["D1"] = "Специальность"
            worksheet["E1"] = "Группа"
            worksheet["F1"] = "База образования (9 или 11 классов)"
            worksheet["G1"] = "Основа образования (бюджет, внебюджет)"
            worksheet["H1"] = "Приказ о зачислении"
            worksheet["I1"] = "Переводной приказ на 2 курс"
            worksheet["J1"] = "Переводной приказ на 3 курс"
            worksheet["K1"] = "Переводной приказ на 4 курс"
            worksheet["L1"] = "Отчисление в связи с окончанием обучения"
            worksheet["M1"] = "Телефон"
            worksheet["N1"] = "Примечание"

            # Высота ячеек
            worksheet.row_dimensions[1].height = 40

            # Ширина ячеек
            worksheet.column_dimensions['A'].width = len(str(len(self.students))) + 3
            worksheet.column_dimensions['B'].width = max([len(i.full_name) for i in self.students]) + 3
            worksheet.column_dimensions['C'].width = max([len(str(i.birth_date)) for i in self.students]) + 3
            worksheet.column_dimensions['D'].width = max([len(i.group.qualification.specialty.code) for i in self.students]) + 3
            worksheet.column_dimensions['E'].width = max([len(str(i.group.full_name)) for i in self.students]) + 3
            worksheet.column_dimensions['F'].width = max([len(i.group.education_base) for i in self.students]) + 3
            worksheet.column_dimensions['G'].width = max([len(i.education_basis) for i in self.students]) + 3
            worksheet.column_dimensions['H'].width = max([len(i.admission_order) for i in self.students]) + 3
            worksheet.column_dimensions['I'].width = max([len(str(i.transfer_to_2nd_year_order)) for i in self.students]) + 3
            worksheet.column_dimensions['J'].width = max([len(str(i.transfer_to_3rd_year_order)) for i in self.students]) + 3
            worksheet.column_dimensions['K'].width = max([len(str(i.transfer_to_4th_year_order)) for i in self.students]) + 3
            worksheet.column_dimensions['L'].width = max([len(str(i.expelled_due_to_graduation)) for i in self.students]) + 3 # это надо переделать
            worksheet.column_dimensions['M'].width = max([len(i.phone) for i in self.students]) + 3

            for i in range(len(self.students)):
                worksheet.cell(row=i+2, column=1).value = i+1
                worksheet.cell(row=i+2, column=2).value = self.students[i].full_name
                worksheet.cell(row=i+2, column=3).value = self.students[i].birth_date
                worksheet.cell(row=i+2, column=4).value = self.students[i].group.qualification.specialty.code
                worksheet.cell(row=i+2, column=5).value = self.students[i].group.full_name
                worksheet.cell(row=i+2, column=6).value = self.students[i].group.education_base
                worksheet.cell(row=i+2, column=7).value = self.students[i].education_basis
                worksheet.cell(row=i+2, column=8).value = self.students[i].admission_order
                worksheet.cell(row=i+2, column=9).value = self.students[i].transfer_to_2nd_year_order
                worksheet.cell(row=i+2, column=10).value = self.students[i].transfer_to_3rd_year_order
                worksheet.cell(row=i+2, column=11).value = self.students[i].transfer_to_4th_year_order
                worksheet.cell(row=i+2, column=12).value = self.students[i].expelled_due_to_graduation
                worksheet.cell(row=i+2, column=13).value = self.students[i].phone
            
            workbook.save(path)

class StatisticsTableGenerator:
    def __init__(self, specialties, qualifications, students) -> None:
        self.specialties = specialties
        self.qualifications = qualifications
        self.students = students

    def generate_document(self, path):
        workbook = Workbook()
        worksheet = workbook.active

        if worksheet is not None:
            from datetime import date
            today = date.today()
            worksheet.title = f"Статистика ({today.day}.{today.month}.{today.year})"
            
            # Заголовки (первый ряд)
            worksheet["A1"] = "№ п/п"
            worksheet["B1"] = "Код"
            worksheet["C1"] = "Направление подготовки/специальности"
            worksheet["D1"] = "Профиль (направленность/ специализация), квалификация для программ СПО"
            worksheet["E1"] = "Форма обучения"
            worksheet["F1"] = "Вид обучения"
            worksheet.merge_cells("G1:R1")
            worksheet["G1"] = "Контингент"
            
            # Заголовки (второй ряд)
            worksheet["A2"] = 1
            worksheet["B2"] = 2
            worksheet["C2"] = 3
            worksheet["D2"] = 4
            worksheet["E2"] = 5
            worksheet["F2"] = 6
            worksheet.merge_cells("G2:J2")
            worksheet["G2"] = 7
            worksheet.merge_cells("K2:M2")
            worksheet["K2"] = 8
            worksheet.merge_cells("N2:Q2")
            worksheet["N2"] = 9
            worksheet["R2"] = 10

            # Заголовки (третий ряд)
            worksheet["A3"] = '-'
            worksheet["B3"] = '-'
            worksheet["C3"] = '-'
            worksheet["D3"] = '-'
            worksheet["E3"] = '-'
            worksheet["F3"] = '-'
            worksheet.merge_cells("G3:J3")
            worksheet["G3"] = "9кл" 
            worksheet.merge_cells("K3:M3")
            worksheet["K3"] = "11кл" 
            worksheet.merge_cells("N3:Q3")
            worksheet["N3"] = "Акад.отпуск"
            worksheet["R3"] = "Итого"

            # Заголовки (четверый ряд)
            worksheet["G4"] = "1 курс"
            worksheet["H4"] = "2 курс"
            worksheet["I4"] = "3 курс"
            worksheet["J4"] = "4 курс"
            worksheet["K4"] = "1 курс"
            worksheet["L4"] = "2 курс"
            worksheet["M4"] = "3 курс"
            worksheet["N4"] = "1 курс"
            worksheet["O4"] = "2 курс"
            worksheet["P4"] = "3 курс"
            worksheet["Q4"] = "4 курс"
            
            entry_count = 1 # порядковый номер квалификации
            entry_index = 5 # индекс последней заполненной строки, вынесен отдельно,
            isip_done = False

            # Проходимся по каждой специальности каждой квалификации
            for specialty in self.specialties:
                qualifications_specialty = self.qualifications.filter(specialty__code=specialty.code)
                for qualification in qualifications_specialty:
                    # Здесь я единожды разбираюсь с 1-м курсом исипа и выделяю место под него в таблице,
                    # после чего единожды заполняю их количество, больше в этом цикле заполнения связанного с ними не будет
                    if specialty.discipline_name == "Информационные системы и программирование" and not isip_done:
                        worksheet.merge_cells(f"G{entry_index}:G{entry_index+(len(qualifications_specialty)-2)}")
                        worksheet.cell(row=entry_index, column=7).fill = PatternFill("solid", fgColor=Color(indexed=3))
                        worksheet.cell(row=entry_index, column=7).value = len(self.students.filter(group__qualification__name="Информационные системы и программирование")
                                                                              .filter(education_basis="Бюджет").filter(is_in_academ=False).filter(is_expelled=False))

                        worksheet.merge_cells(f"G{entry_index+(len(qualifications_specialty))-1}:G{entry_index+(len(qualifications_specialty)*2) - 3}")
                        worksheet.cell(row=entry_index+len(qualifications_specialty)-1, column=7).fill = PatternFill("solid", fgColor=Color(indexed=7))
                        worksheet.cell(row=entry_index+len(qualifications_specialty)-1, column=7).value = len(self.students.filter(group__qualification__name="Информационные системы и программирование")
                                                                                                              .filter(education_basis="Внебюджет").filter(is_in_academ=False).filter(is_expelled=False))

                        # Академ отпуск
                        worksheet.merge_cells(f"N{entry_index}:N{entry_index+(len(qualifications_specialty)-2)}")
                        worksheet.cell(row=entry_index, column=14).fill = PatternFill("solid", fgColor=Color(indexed=3))
                        worksheet.cell(row=entry_index, column=14).value = len(self.students.filter(group__qualification__name="Информационные системы и программирование")
                                                                               .filter(education_basis="Бюджет").filter(is_in_academ=True).filter(is_expelled=False))

                        worksheet.merge_cells(f"N{entry_index+len(qualifications_specialty) - 1}:N{entry_index+(len(qualifications_specialty)*2 - 3)}")
                        worksheet.cell(row=entry_index+len(qualifications_specialty)-1, column=14).fill = PatternFill("solid", fgColor=Color(indexed=7))
                        worksheet.cell(row=entry_index+len(qualifications_specialty)-1, column=14).value = len(self.students.filter(group__qualification__name="Информационные системы и программирование")
                                                                                                               .filter(education_basis="Внебюджет").filter(is_in_academ=True).filter(is_expelled=False))

                        worksheet.merge_cells(f"S{entry_index}:S{entry_index+(len(qualifications_specialty)*2 - 3)}")
                        worksheet.cell(row=entry_index, column=19).value = "Итог без учета 1-го курса специальности 09.02.07"

                        isip_done = True
                    
                    # В случае если нам попадается эта квалификация (это 1-ый курс исипа),
                    # ничего не делаем, т.к. до этого мы уже с ними разобрались
                    if qualification.discipline_name == "Информационные системы и программирование":
                        continue
                    
                    worksheet.cell(row=entry_index,column=1).value = entry_count
                    worksheet.cell(row=entry_index,column=2).value = specialty.code
                    worksheet.cell(row=entry_index,column=3).value = specialty.discipline_name
                    worksheet.cell(row=entry_index,column=4).value = qualification.discipline_name
                    worksheet.cell(row=entry_index,column=5).value = "Очная"

                    students_qualification = self.students.filter(group__qualification__name=qualification.discipline_name)
                    worksheet.merge_cells(f"A{entry_index}:A{entry_index+1}")
                    worksheet.merge_cells(f"B{entry_index}:B{entry_index+1}")
                    worksheet.merge_cells(f"C{entry_index}:C{entry_index+1}")
                    worksheet.merge_cells(f"D{entry_index}:D{entry_index+1}")
                    worksheet.merge_cells(f"E{entry_index}:E{entry_index+1}")
                    
                    worksheet.cell(row=entry_index, column=6).value = "Бюджет"
                    worksheet.cell(row=entry_index+1, column=6).value = "Договор"

                    # Суммарное количество всех студентов на квалификации (Бюджет и договор)
                    worksheet.cell(row=entry_index, column=18).value = len(students_qualification.filter(education_basis="Бюджет"))
                    worksheet.cell(row=entry_index+1, column=18).value = len(students_qualification.filter(education_basis="Внебюджет"))

                    # 9 классники
                    for course in range(1, 5):
                        if not(course == 1 and specialty.discipline_name == "Информационные системы и программирование"):
                            worksheet.cell(row=entry_index, column=6+course).value = len(students_qualification.filter(group__current_course=course).filter(education_basis="Бюджет")
                                                                                         .filter(education_base="Основное общее").filter(is_in_academ=False).filter(is_expelled=False))
                            worksheet.cell(row=entry_index+1, column=6+course).value = len(students_qualification.filter(group__current_course=course).filter(education_basis="Внебюджет")
                                                                                           .filter(education_base="Основное общее").filter(is_in_academ=False).filter(is_expelled=False))
                            # Академ отпуска (и для 9 и для 11)
                            worksheet.cell(row=entry_index, column=13+course).value = len(students_qualification.filter(group__current_course=course).filter(education_basis="Бюджет")
                                                                                          .filter(is_in_academ=True).filter(is_expelled=False))
                            worksheet.cell(row=entry_index+1, column=13+course).value = len(students_qualification.filter(group__current_course=course).filter(education_basis="Внебюджет")
                                                                                          .filter(is_in_academ=True).filter(is_expelled=False))
                    # 11 классники
                    for course in range(1, 4):
                         worksheet.cell(row=entry_index, column=10+course).value = len(students_qualification.filter(group__current_course=course).filter(education_basis="Бюджет")
                                                                                         .filter(education_base="Среднее общее").filter(is_in_academ=False).filter(is_expelled=False))
                         worksheet.cell(row=entry_index+1, column=10+course).value = len(students_qualification.filter(group__current_course=course).filter(education_basis="Внебюджет")
                                                                                         .filter(education_base="Среднее общее").filter(is_in_academ=False).filter(is_expelled=False))
                    entry_index += 2

                    # elif len(students_qualification) == 0:
                    #     continue

                    entry_count += 1
            
            # Итог под каждый курс
            for column_num in range(7, 19):
                worksheet.cell(row=entry_index, column=column_num).fill = PatternFill("solid", fgColor=Color(indexed=5))

            worksheet.cell(row=entry_index, column=7).value = len(self.students.filter(education_base="Основное общее").filter(group__current_course=1)
                                                                   .filter(is_in_academ=False).filter(is_expelled=False))
            worksheet.cell(row=entry_index, column=8).value = len(self.students.filter(education_base="Основное общее").filter(group__current_course=2)
                                                                   .filter(is_in_academ=False).filter(is_expelled=False))
            worksheet.cell(row=entry_index, column=9).value = len(self.students.filter(education_base="Основное общее").filter(group__current_course=3)
                                                                   .filter(is_in_academ=False).filter(is_expelled=False))
            worksheet.cell(row=entry_index, column=10).value = len(self.students.filter(education_base="Основное общее").filter(group__current_course=4)
                                                                   .filter(is_in_academ=False).filter(is_expelled=False))
            worksheet.cell(row=entry_index, column=11).value = len(self.students.filter(education_base="Среднее общее").filter(group__current_course=1)
                                                                   .filter(is_in_academ=False).filter(is_expelled=False))
            worksheet.cell(row=entry_index, column=12).value = len(self.students.filter(education_base="Среднее общее").filter(group__current_course=2)
                                                                   .filter(is_in_academ=False).filter(is_expelled=False))
            worksheet.cell(row=entry_index, column=13).value = len(self.students.filter(education_base="Среднее общее").filter(group__current_course=3)
                                                                   .filter(is_in_academ=False).filter(is_expelled=False))
            worksheet.cell(row=entry_index, column=14).value = len(self.students.filter(group__current_course=1).filter(is_in_academ=True).filter(is_expelled=False))
            worksheet.cell(row=entry_index, column=15).value = len(self.students.filter(group__current_course=2).filter(is_in_academ=True).filter(is_expelled=False))
            worksheet.cell(row=entry_index, column=16).value = len(self.students.filter(group__current_course=3).filter(is_in_academ=True).filter(is_expelled=False))
            worksheet.cell(row=entry_index, column=17).value = len(self.students.filter(group__current_course=4).filter(is_in_academ=True).filter(is_expelled=False))

            # Итог со всеми студентами
            worksheet.cell(row=entry_index, column=18).value = len(self.students.filter(is_expelled=False))
            worksheet.cell(row=entry_index+2, column=17).value = "Итого:"
            worksheet.cell(row=entry_index+2, column=18).value = len(self.students.filter(is_expelled=False))
            
            # Итог с ИСИП
            worksheet.cell(row=entry_index+5, column=2).fill = PatternFill("solid", fgColor=Color(indexed=3))
            worksheet.cell(row=entry_index+5, column=2).value = len(self.students.filter(group__qualification__name="Информационные системы и программирование")
                                                                  .filter(education_basis="Бюджет").filter(is_in_academ=False).filter(is_expelled=False))
            worksheet.cell(row=entry_index+6, column=2).fill = PatternFill("solid", fgColor=Color(indexed=7))
            worksheet.cell(row=entry_index+6, column=2).value = len(self.students.filter(group__qualification__name="Информационные системы и программирование")
                                                                  .filter(education_basis="Внебюджет").filter(is_in_academ=False).filter(is_expelled=False))
            worksheet.cell(row=entry_index+5, column=3).value = "- бюджет"
            worksheet.cell(row=entry_index+6, column=3).value = "- договор"
        workbook.save(path)


class VacationTableGenerator:
    def __init__(self, students):
        self.students = students

    def generate_document(self, path):
        workbook = Workbook()
        worksheet = workbook.active
        if worksheet is not None:
            # Заголовок
            worksheet["A1"].value = "№"
            worksheet["B1"].value = "ФИО"
            worksheet["C1"].value = "Дата рождения"           
            worksheet["D1"].value = "Специальность"
            worksheet["E1"].value = "Курс, с которого ушел"
            worksheet["F1"].value = "Дата выхода"
            worksheet["G1"].value = "Период"
            worksheet["H1"].value = "Причина"
            worksheet["I1"].value = "Группа"
            worksheet["J1"].value = "База образования (9 или 11 классов)"
            worksheet["K1"].value = "Основа образования (бюджет, внебюджет)"
            worksheet["L1"].value = "Приказ о зачислении"
            worksheet["M1"].value = "Переводной приказ на 2 курс"
            worksheet["N1"].value = "Переводной приказ на 3 курс"
            worksheet["O1"].value = "Переводной приказ на 4 курс"
            worksheet["P1"].value = "Телефон"
            worksheet["Q1"].value = "Примечание"
            
            entry_index = 2
            for student in self.students:
                if student.is_in_academ and not student.is_expelled:
                    worksheet.cell(row=entry_index, column=1).value = entry_index - 1
                    worksheet.cell(row=entry_index, column=2).value = student.full_name
                    worksheet.cell(row=entry_index, column=3).value = student.birth_date
                    worksheet.cell(row=entry_index, column=4).value = student.group.qualification.specialty.code
                    worksheet.cell(row=entry_index, column=5).value = student.left_course
                    worksheet.cell(row=entry_index, column=6).value = student.academ_return_date
                    worksheet.cell(row=entry_index, column=7).value = f"{student.academ_leave_date} по {student.academ_return_date}"
                    worksheet.cell(row=entry_index, column=8).value = student.reason_of_academ
                    worksheet.cell(row=entry_index, column=9).value = student.group.full_name
                    worksheet.cell(row=entry_index, column=10).value = student.education_base
                    worksheet.cell(row=entry_index, column=11).value = student.education_basis
                    worksheet.cell(row=entry_index, column=12).value = student.admission_order
                    worksheet.cell(row=entry_index, column=13).value = student.transfer_to_2nd_year_order
                    worksheet.cell(row=entry_index, column=14).value = student.transfer_to_3rd_year_order
                    worksheet.cell(row=entry_index, column=15).value = student.transfer_to_4th_year_order
                    worksheet.cell(row=entry_index, column=16).value = student.phone

                    entry_index += 1
        workbook.save(path)


class MovementTableGenerator:
    def __init__(self, movements):
        self.movements = movements

    def generate_document(self, path):
        workbook = Workbook()
        worksheet = workbook.active

        if worksheet is not None:
            # Заголовки
            worksheet["A1"] = "Номер приказа"
            worksheet["B1"] = "Тип действия"
            worksheet["C1"] = "Дата действия"
            worksheet["D1"] = "ФИО студента"
            worksheet["E1"] = "Предыдущая группа"
            worksheet["F1"] = "Новая группа"
            worksheet["G1"] = "Специальность"
            worksheet["H1"] = "База образования (9 или 11 классов)"
            worksheet["I1"] = "Основа образования (бюджет или внебюджет)"

#            worksheet["J1"] = "Приказ о зачислении"
#           worksheet["K1"] = "Переводной приказ на 2 курс"
#            worksheet["L1"] = "Переводной приказ на 3 курс"
#            worksheet["M1"] = "Переводной приказ на 4 курс"
#            worksheet["N1"] = "Телефон"
 #           worksheet["O1"] = "Примечание"

#            entry_index = 2

#            for student in students:
#                if student.is_expelled and not student.expelled_due_to_graduation:
#                    worksheet.cell(row=entry_index, column=1).value = student.expell_order
#                    worksheet.cell(row=entry_index, column=2).value = student.date_of_expelling
#                    worksheet.cell(row=entry_index, column=3).value = student.birth_date
 #                   worksheet.cell(row=entry_index, column=4).value = student.reason_of_expelling
 #                   worksheet.cell(row=entry_index, column=5).value = student.birth_date
 #                   worksheet.cell(row=entry_index, column=6).value = student.group.qualification.specialty.name
#                    worksheet.cell(row=entry_index, column=7).value = student.group.full_name
#                    worksheet.cell(row=entry_index, column=8).value = student.education_base
#                    worksheet.cell(row=entry_index, column=9).value = student.education_basis
#                    worksheet.cell(row=entry_index, column=10).value = student.admission_order
 #                   worksheet.cell(row=entry_index, column=11).value = student.transfer_to_2nd_year_order
#                    worksheet.cell(row=entry_index, column=12).value = student.transfer_to_3rd_year_order
#                    worksheet.cell(row=entry_index, column=13).value = student.transfer_to_4th_year_order
#                    worksheet.cell(row=entry_index, column=14).value = student.phone
#
 #       workbook.save(path)
            # Женин и мой вариант. Надо переделать шоб нормально было
            worksheet["J1"] = "Телефон"

            # Заполнение данными из ContingentMovement
            for i, movement in enumerate(self.movements, start=2):
                worksheet.cell(row=i, column=1).value = movement.order_number
                worksheet.cell(row=i, column=2).value = movement.get_action_type_display()  # Читаемое название типа действия
                worksheet.cell(row=i, column=3).value = str(movement.action_date)
                worksheet.cell(row=i, column=4).value = movement.student.full_name
                worksheet.cell(row=i, column=5).value = movement.previous_group
                worksheet.cell(row=i, column=6).value = movement.new_group
                worksheet.cell(row=i, column=7).value = movement.student.group.qualification.specialty.code
                worksheet.cell(row=i, column=8).value = movement.student.group.education_base
                worksheet.cell(row=i, column=9).value = movement.student.education_basis
                worksheet.cell(row=i, column=10).value = movement.student.phone

            workbook.save(path)

class GradebookDocumentGenerator:
    def __init__(self, gradebook) -> None:
        self.gradebook = gradebook
        self.gradebook_name = gradebook.name
        self.gradebook_number = gradebook.number
        self.gradebook_discipline = gradebook.discipline
        self.gradebook_semester = gradebook.semester_number
        self.gradebook_course = gradebook.group.current_course
        self.gradebook_group = gradebook.group.full_name
        self.gradebook_students = gradebook.students
        self.gradebook_specialty = gradebook.group.qualification.specialty
        self.gradebook_teachers = [str(i[5]) for i in list(gradebook.teachers.values_list())]

    def generate_document(self, path) -> None:
        document = Document()
        sections = document.sections
        for section in sections:
            section.page_width = Cm(21.0)
            section.page_height = Cm(29.70)
            section.left_margin = Cm(3.0)
            section.right_margin = Cm(1.49)
            section.top_margin = Cm(1.50)
            section.bottom_margin = Cm(2.0)

        styles = document.styles
        # Определенние стилей:
        # Heading 1
        heading_style = styles.add_style('Our_Heading 1', WD_STYLE_TYPE.PARAGRAPH)
        heading_style.base_style = styles['Heading 1']
        heading_style.font.name = 'Times New Roman'
        heading_style.font.size = Pt(12)
        heading_style.font.color.rgb = RGBColor(0, 0, 0)

        # Normal
        normal_style = styles.add_style('Our_Normal', WD_STYLE_TYPE.PARAGRAPH)
        normal_style.base_style = styles['Normal']
        normal_style.font.name = 'Times New Roman'
        normal_style.font.size = Pt(10)

        # Subtitle
        subtitle_style = styles.add_style('Our_subtitle', WD_STYLE_TYPE.PARAGRAPH)
        subtitle_style.base_style = styles['Our_subtitle']
        subtitle_style.font.name = 'Times New Roman'
        subtitle_style.font.size = Pt(14)
        subtitle_style.font.bold = True

        # Отсюда идет форматирование самого документа
        # Заголовок в начале документа
        paragraph = document.add_paragraph('МИНИСТЕРСТВО НАУКИ И ВЫСШЕГО ОБРАЗОВАНИЯ РОССИЙСКОЙ ФЕДЕРАЦИИ\n', 'Our_Normal')
        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.style.font.bold = True
        paragraph.style.font.size = Pt(9)
        run = paragraph.add_run('''федеральное государственное бюджетное образовательное учреждение высшего образования 
        «Российский экономический университет имени Г.В. Плеханова»''')
        run.bold = True
        run.font.size = Pt(10)
        paragraph.paragraph_format.space_after = 0

        paragraph = document.add_heading('Московский приборостроительный техникум\n', 1)
        paragraph.style = 'Our_Heading 1'
        paragraph.paragraph_format.space_before = 0
        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Subtitle с номером ведомости
        paragraph = document.add_paragraph(f'{self.gradebook_name.upper()}   ', 'Our_subtitle')
        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run(f'№ {self.gradebook_number}').underline = True

        # Параграф с описанием про кого и от кого ведомость (не закончено: нужно разобраться со стилями разных параграфов
        # В случае если курсовой проект надо писать не "по дисциплине", а "курсовой проект"
        if self.gradebook_name == 'Ведомость защиты курсового проекта':
            paragraph = document.add_paragraph('Курсовой проект: ', 'Our_Normal')
        else:
            paragraph = document.add_paragraph('По дисциплине/МДК: ', 'Our_Normal')

        paragraph.style.font.bold = False
        paragraph.add_run(f'{self.gradebook_discipline}\n').font.size = Pt(12)
        paragraph.add_run(str(self.gradebook_course)).bold = True
        paragraph.add_run(' курс ').font.size = Pt(10)
        paragraph.add_run(str(self.gradebook_semester)).bold = True
        paragraph.add_run(' семестр ').font.size = Pt(10)
        paragraph.add_run(self.gradebook_group).bold = True
        paragraph.add_run(' группа\n').font.size = Pt(10)
        paragraph.add_run('Специальность: ').font.size = Pt(10)
        paragraph.add_run(f'{self.gradebook_specialty.code} "{self.gradebook_specialty.name}"\n').bold = True
        paragraph.add_run('Форма обучения: ').font.size = Pt(10)
        paragraph.add_run('Очная\n').bold = True
        paragraph.add_run('Преподаватель: ').font.size = Pt(10)
        paragraph.add_run(' '.join(self.gradebook_teachers))
        paragraph.add_run('\n')
        
        # Сделать отсутствие колонки с билетами в случае если курсовой или успеваемость
        
        if self.gradebook_name == 'Ведомость успеваемости' or self.gradebook_name == 'Ведомость защиты курсового проекта':
            table = document.add_table(rows=1, cols=4)
            # Устанавливаем ширину столбцов (в дюймах)
            column_widths = [
                Inches(0.2),  # № П/П
                Inches(10),  # ФИО
                Inches(0.5),  # Оценка
                Inches(0.5)  # Подпись
            ]

            for i, width in enumerate(column_widths):
                table.columns[i].width = width

            table.style = 'Table Grid'
            table.style.font.name = "Times New Roman"
            table.style.alignment = WD_ALIGN_PARAGRAPH.CENTER
            table.style.font.size = Pt(11)
            # Заголовки таблицы
            header_cells = table.rows[0].cells
            header_cells[0].text = '№ П/П'
            header_cells[1].text = 'Фамилия, имя, отчество студента'
            header_cells[2].text = 'Оценка'
            header_cells[3].text = 'Подпись экзаменатора'

            # Выравнивание текста в заголовках
            for cell in header_cells:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            counter = 1
            for gradebook_student in GradebookStudents.objects.filter(gradebook=self.gradebook):
                row_cells = table.add_row().cells
                row_cells[0].text = str(counter)
                row_cells[1].text = gradebook_student.student.full_name
                row_cells[2].text = str(gradebook_student.grade)

                # Выравнивание для ФИО
                row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

                # Выравнивание для числовых колонок
                for i in [0, 2, 3]:  # Колонки с числами
                    row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    row_cells[i].paragraphs[0].alignment = WD_ALIGN_VERTICAL.CENTER

                counter += 1
        else:
            table = document.add_table(rows=1, cols=5)
            # Устанавливаем ширину столбцов (в дюймах)
            column_widths = [
                Inches(0.2),  # № П/П
                Inches(0.2),  # № экз.билета
                Inches(10),  # ФИО
                Inches(0.5),  # Оценка
                Inches(0.5)  # Подпись
            ]

            for i, width in enumerate(column_widths):
                table.columns[i].width = width

            table.style = 'Table Grid'
            table.style.font.name = "Times New Roman"
            table.style.font.size = Pt(11)
            # Заголовки таблицы
            header_cells = table.rows[0].cells
            header_cells[0].text = '№ П/П'
            header_cells[1].text = '№ экз.билета'
            header_cells[2].text = 'Фамилия, имя, отчество студента'
            header_cells[3].text = 'Оценка'
            header_cells[4].text = 'Подпись экзаменатора'

            # Выравнивание текста в заголовках
            for cell in header_cells:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            counter = 1
            for gradebook_student in GradebookStudents.objects.filter(gradebook=self.gradebook):
                row_cells = table.add_row().cells
                row_cells[0].text = str(counter)
                row_cells[1].text = str(gradebook_student.ticket_number)
                row_cells[2].text = gradebook_student.student.full_name
                row_cells[3].text = str(gradebook_student.grade)

                # Выравнивание для ФИО
                row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

                # Выравнивание для числовых колонок
                for i in [0, 1, 3]:  # Колонки с числами
                    row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    row_cells[i].paragraphs[0].alignment = WD_ALIGN_VERTICAL.CENTER
                counter += 1

        paragraph = document.add_paragraph('\n')

        footer_table = document.add_table(rows=1, cols=3)
        footer_table.autofit = False  # Отключаем авто-подгонку
        footer_table.rows[0].cells[1].width = Cm(8)
        footer_table.rows[0].cells[2].width = Cm(4)

        footer_table.rows[0].cells[0].text = '«      »         ' + str(self.gradebook.date_of_closing.year) + ' года'


        footer_table.rows[0].cells[1].text = 'Подпись преподавателя:'
        for paragraph in footer_table.rows[0].cells[1].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT  # Выравнивание по правому краю

        footer_table.rows[0].cells[2].text = '_______________'
        for paragraph in footer_table.rows[0].cells[2].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT  # Выравнивание по правому краю

        # Устанавливаем стиль для ВСЕХ ячеек таблицы
        for row in footer_table.rows:
            for cell in row.cells:
                # Настройка шрифта для каждого параграфа в ячейке
                for paragraph in cell.paragraphs:
                    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(10)


        paragraph = document.add_paragraph('\n')

        # Table for grades count
        footer_table2 = document.add_table(rows=6, cols=6)

        footer_table2.style.font.name = "Times New Roman"
        footer_table2.autofit = False  # Отключаем авто-подгонку
        footer_table2.rows[0].cells[4].width = Cm(4)

        footer_table2.rows[0].cells[4].text = 'Всего оценок: '
        for paragraph in footer_table2.rows[0].cells[4].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT  # Выравнивание по правому краю
        footer_table2.rows[0].cells[5].text = str(GradebookStudents.objects.filter(gradebook=self.gradebook).count())

        footer_table2.rows[1].cells[4].text = 'в том числе «5» –'
        for paragraph in footer_table2.rows[1].cells[4].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT  # Выравнивание по правому краю
        footer_table2.rows[1].cells[5].text = str(GradebookStudents.objects.filter(gradebook=self.gradebook, grade="Отлично").count())

        footer_table2.rows[2].cells[4].text = '«4» –'
        for paragraph in footer_table2.rows[2].cells[4].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT  # Выравнивание по правому краю
        footer_table2.rows[2].cells[5].text = str(GradebookStudents.objects.filter(gradebook=self.gradebook, grade="Хорошо").count())

        footer_table2.rows[3].cells[4].text = '«3» –'
        for paragraph in footer_table2.rows[3].cells[4].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT  # Выравнивание по правому краю
        footer_table2.rows[3].cells[5].text = str(GradebookStudents.objects.filter(gradebook=self.gradebook, grade="Удовлетворительно").count())

        footer_table2.rows[4].cells[4].text = '«2» –'
        for paragraph in footer_table2.rows[4].cells[4].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT  # Выравнивание по правому краю
        footer_table2.rows[4].cells[5].text = str(GradebookStudents.objects.filter(gradebook=self.gradebook, grade="Неудовлетворительно").count())

        footer_table2.rows[5].cells[4].text = '«н/я» –'
        for paragraph in footer_table2.rows[5].cells[4].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT # Выравнивание по правому краю
        footer_table2.rows[5].cells[5].text = str(GradebookStudents.objects.filter(gradebook=self.gradebook, grade="Неявка").count())

        # Устанавливаем стиль для ВСЕХ ячеек таблицы
        for row in footer_table2.rows:
            for cell in row.cells:
                # Настройка шрифта для каждого параграфа в ячейке
                for paragraph in cell.paragraphs:
                    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(10)

        document.save(path)

# def tableStudentsParse(path):
#     ext = os.path.splitext(path)[1].lower()
#     data = []
#     if ext == ".xlsx":
#         workbook = load_workbook(filename=path, data_only=True)
#         worksheet = workbook.active
#         if worksheet is not None:
#             for row in worksheet.iter_rows(values_only=True):
#                 data.append(list(row))
#     elif ext == ".xls":
#         import xlrd
#         workbook = xlrd.open_workbook(path)
#         worksheet = workbook.sheet_by_index(0)
#         for row_index in range(worksheet.nrows):
#             data.append(worksheet.row_values(row_index))
#     else:
#         ValueError("Неподдерживаемый формат файла")
#     
#     students = []
#     for student_row in data:
#        student = Student(full_name=student_row[1], )

if __name__ == "__main__":
    # students = Student.objects.all()
    # specialties = Specialty.objects.all()
    # qualifications = Qualification.objects.all()
    # groups = GroupStudents.objects.all()
    # movements = ContingentMovement.objects.all()
    #
    # gtg = GroupTableGenerator(groups=groups)
    # ctg = CourseTableGenerator(students=students, education_base="Основное общее", course="2")
    # stg = StatisticsTableGenerator(specialties=specialties, qualifications=qualifications, students=students)
    # vtg = VacationTableGenerator(students=students)
    # mtg = MovementTableGenerator(movements=movements)
    #
    # gtg.generate_document("test_groups.xlsx")
    # ctg.generate_document("test_course.xlsx")
    # stg.generate_document("test_statistics.xlsx")
    # vtg.generate_document("test_vacation.xlsx")
    # mtg.generate_document("test_movement.xlsx")
    #

    gradebook_exam = Gradebook.objects.get(number="ааа1123")
    gradebook_courseProject = Gradebook.objects.get(number="546NN")

    gdg = GradebookDocumentGenerator(gradebook_exam)

    gdg.generate_document("test_gradebook_exam.docx")

    gdg = GradebookDocumentGenerator(gradebook_courseProject)
    gdg.generate_document("test_gradebook_courseProject.docx")

