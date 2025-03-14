import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def replace_text_in_paragraph(paragraph, key, value):
    if key in paragraph.text:
        inline = paragraph.runs
        # Loop added to work with runs (strings with same style)
        for i in range(len(inline)):
            if key in inline[i].text:
                text = inline[i].text.replace(key, value)
                inline[i].text = text


def replace_text_in_table(table, key, value):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                replace_text_in_paragraph(paragraph, key, value)


def create_document_from_template(template_path, output_path, data):
    document = docx.Document(template_path)

    for key, value in data.items():
        for paragraph in document.paragraphs:
            replace_text_in_paragraph(paragraph, key, value)
        for table in document.tables:
            replace_text_in_table(table, key, value)

    document.save(output_path)


def generate_report(data, template_path, output_path):
    try:
        doc = docx.Document(template_path)
        for key, value in data.items():
            for paragraph in doc.paragraphs:
                if key in paragraph.text:
                    paragraph.text = paragraph.text.replace(key, str(value))
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            if key in paragraph.text:
                                paragraph.text = paragraph.text.replace(key, str(value))
        doc.save(output_path)
        return True
    except Exception as e:
        print(f"Error generating report: {e}")
        return False